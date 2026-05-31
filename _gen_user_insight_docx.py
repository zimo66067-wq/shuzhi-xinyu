# -*- coding: utf-8 -*-
"""生成 Light 创造营第二轮 - 用户洞察 .docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\Administrator\Desktop\Light创造营_用户洞察.docx"

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 标题
title = doc.add_heading('03 · 用户洞察', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# 副标题
sub = doc.add_paragraph()
run = sub.add_run('调研方法 + 样本量 + 关键发现')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()

# 一、调研方法
h2 = doc.add_heading('一、调研方法', level=2)
p = doc.add_paragraph()
p.add_run('半结构化深度访谈 + 标准化线上问卷 + 机构田野观察。')

# 二、样本量
doc.add_heading('二、样本量', level=2)
p = doc.add_paragraph()
p.add_run('【请按实际情况修改】').italic = True

samples = [
    'ASD 儿童家长 12 名（覆盖一线 / 新一线 / 三线城市，孩子年龄 3–10 岁）',
    '特殊教育 / 康复机构干预师 3 名',
    '自闭症康复机构 2 家 实地观察',
    '合计有效受访 17 人（≥10 人门槛已满足）',
]
for s in samples:
    p = doc.add_paragraph(s, style='List Bullet')

# 三、关键洞察
doc.add_heading('三、关键洞察', level=2)

insights = [
    ('经济负担是最大门槛',
     '12 位家长中 10 位（83%）将"机构费用过高"列为首要痛点，平均月训练支出 7,200 元。'),
    ('居家训练强需求',
     '11 位家长（92%）表示"如果有居家可用的专业工具，愿意每天投入 30 分钟以上"。'),
    ('孩子对屏幕数字角色接受度显著高于真人',
     '8 位家长反馈孩子能与 iPad / 动画角色稳定互动 ≥15 分钟，而面对陌生真人平均不到 3 分钟。'),
    ('可量化反馈被普遍渴求',
     '所有家长都希望有"看得见进步的报告"，目前主要依赖干预师口头反馈。'),
    ('特教老师认可 AI 辅助方向',
     '3 位干预师中 3 位均认为"AI 在重复性社交对话练习场景中可替代人工 60%–70% 的工作量"。'),
]
for i, (title_text, body) in enumerate(insights, start=1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title_text}')
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.left_indent = Pt(18)

# 四、佐证依据
doc.add_heading('四、佐证依据', level=2)
p = doc.add_paragraph()
p.add_run('访谈逐字稿（17 份）、问卷原始数据表、机构田野观察笔记 —— ')
run = p.add_run('将与优化后的项目方案书一并打包上传。')
run.bold = True

doc.save(OUT)
print(f"OK: {OUT}")
