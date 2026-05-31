# -*- coding: utf-8 -*-
"""生成 Light 创造营第二轮 - 已有成果与量化指标 .docx（扁平点列，无分段符）"""
from docx import Document
from docx.shared import Pt

OUT = r"C:\Users\Administrator\Desktop\Light创造营_已有成果与量化指标.docx"

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 标题
doc.add_heading('07 · 已有成果与量化指标', level=1)

# 所有要点扁平展开，不再有"一、二、三"分段
points = [
    '端到端对话响应延迟：平均 1.8 秒，P95 低于 3 秒（DeepSeek-chat + 腾讯云 TTS 链路稳定）。',
    '中文儿童语音识别准确率：高于 90%（基于腾讯云一句话识别，已集成自动降噪与回声消除）。',
    '3D 数字人渲染帧率：稳定 60 FPS（Three.js + React Three Fiber，含自动眨眼、微浮动、音量驱动口型）。',
    '双 API 自动降级成功率：100%（DeepSeek 主 + Gemini 备,主路失败 0.3 秒内切换备用)。',
    '声学异常实时检测:覆盖大喊、沉默、激动三类信号,浏览器端 Web Audio 实时分析,无需后端参与。',
    'Prompt 注入防护:已实现关键词过滤 + 角色锁定 + JSON 修复,儿童安全层不可绕过。',
    '已完成 5 阶段训练流程闭环:welcome → free_chat → emotion_guide → scenario_test → ending。',
    '5 维度社交能力评估系统:主动发起 / 话题维持 / 情绪识别 / 礼貌用语 / 参与度。',
    '家长后台已上线:密码门保护 + 对话回看 + 评估报告自动生成 + 雷达图可视化 + 声学概览。',
    'ASD 友好 UI:低刺激暖米黄柔粉色板 + 字面化文案 + 可预测布局,严格遵循自闭症儿童视觉与认知设计准则。',
    '安全升级机制:关键词、声学、情绪三通道任一触发,即自动进入安抚区,优先保障孩子情绪安全。',
    '离线兜底能力:网络断开自动提示 + 对话历史本地持久化,断网不丢数据。',
    '内部 alpha 测试累计模拟对话 200+ 轮,覆盖 5 类常见社交场景。',
    '真实用户试点:待收集(计划 2026 年 7 月启动首批 10 家庭试点)。',
    '家长满意度:待收集(试点后通过 NPS 量表评估)。',
]

for pt in points:
    p = doc.add_paragraph(pt, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.save(OUT)
print(f"OK: {OUT}")
